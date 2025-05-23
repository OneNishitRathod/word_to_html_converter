import mammoth
import html
import re
import os
import base64
from PIL import Image
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import zipfile

def convert_docx_to_html(docx_path, font_family=None, font_size=None, max_image_width=800):
    """
    Convert a Word document to HTML using mammoth with enhanced content handling.
    
    Args:
        docx_path (str): Path to the Word document file.
        font_family (str, optional): Font family to apply to the HTML content.
        font_size (str, optional): Font size to apply to the HTML content.
        max_image_width (int): Maximum width for images in pixels.
        
    Returns:
        tuple: (html_content, raw_html_content)
            - html_content: HTML content for code display
            - raw_html_content: HTML content for rendering
    """
    
    # Image conversion function
    def convert_image(image):
        try:
            with image.open() as image_bytes:
                # Convert to PIL Image
                image_data = image_bytes.read()
                with BytesIO(image_data) as bio:
                    pil_image = Image.open(bio)
                    # Ensure image is loaded
                    pil_image.load()
                    
                    # Convert TIFF to PNG for better browser support
                    if pil_image.format.upper() == 'TIFF':
                        new_image = Image.new('RGB', pil_image.size, (255, 255, 255))
                        new_image.paste(pil_image)
                        pil_image = new_image
                    
                    # Get original format or default to PNG
                    image_format = pil_image.format or 'PNG'
                    mime_type = f"image/{image_format.lower()}"
                    
                    # Resize if necessary while maintaining aspect ratio
                    if pil_image.width > max_image_width:
                        ratio = max_image_width / pil_image.width
                        new_height = int(pil_image.height * ratio)
                        pil_image = pil_image.resize((max_image_width, new_height), Image.Resampling.LANCZOS)
                    
                    # Save to buffer with optimizations
                    buffer = BytesIO()
                    if image_format.upper() == 'JPEG':
                        pil_image.save(buffer, format=image_format, optimize=True, quality=85)
                    else:
                        pil_image.save(buffer, format='PNG', optimize=True)
                    
                    # Convert to base64
                    img_str = base64.b64encode(buffer.getvalue()).decode()
                    
                    # Return properly formatted data URL
                    return {
                        "src": f"data:image/png;base64,{img_str}",
                        "style": "max-width: 100%; height: auto; display: block; margin: 10px auto;"
                    }
        except Exception as e:
            print(f"Error processing image: {str(e)}")
            return {"src": "", "alt": "Image conversion failed"}

    # Enhanced style map for better list handling
    style_map = """
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Heading 4'] => h4:fresh
        p[style-name='Heading 5'] => h5:fresh
        p[style-name='Heading 6'] => h6:fresh
        p[style-name='Title'] => h1.title:fresh
        p[style-name='Subtitle'] => h2.subtitle:fresh
        p[style-name='Quote'] => blockquote:fresh
        p[style-name='Intense Quote'] => blockquote.intense:fresh
        table => table.document-table
        tr => tr
        td => td
        p[style-name='List Paragraph'] => li.list-item:fresh
        p[style-name='List Bullet'] => li.list-bullet:fresh
        p[style-name='List Bullet 2'] => li.list-bullet-2:fresh
        p[style-name='List Bullet 3'] => li.list-bullet-3:fresh
        p[style-name='List Number'] => li.list-number:fresh
        p[style-name='List Number 2'] => li.list-number-2:fresh
        p[style-name='List Number 3'] => li.list-number-3:fresh
        r[style-name='List Bullet'] => span.list-bullet
        r[style-name='List Number'] => span.list-number
    """

    # Convert document with image handling
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(
            docx_file,
            style_map=style_map,
            convert_image=mammoth.images.img_element(convert_image)
        )
        
        html_content = result.value
        
        # Post-process HTML for better formatting
        html_content = enhance_html_formatting(html_content)
        
        # Apply custom font styling if specified
        if font_family or font_size:
            style_attrs = []
            if font_family:
                style_attrs.append(f"font-family: {font_family}")
            if font_size:
                style_attrs.append(f"font-size: {font_size}")
            
            style_str = "; ".join(style_attrs)
            html_content = f'<div style="{style_str}">{html_content}</div>'
        
        # Add comprehensive CSS for better rendering
        full_html = create_complete_html(html_content)
        
        return html_content, full_html

def enhance_html_formatting(html_content):
    """
    Enhance HTML formatting for better display of complex content.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Convert flat list items into properly nested structure
    def nest_list_items(items):
        if not items:
            return None
            
        # Create the main list
        first_item = items[0]
        is_bullet = any('bullet' in c for c in first_item.get('class', []))
        main_list = soup.new_tag('ul' if is_bullet else 'ol')
        
        current_level = 1
        current_parent = main_list
        parents_stack = [(1, main_list)]
        
        for item in items:
            # Determine item level
            item_class = item.get('class', [])
            item_level = 1
            is_current_bullet = any('bullet' in c for c in item_class)
            
            for c in item_class:
                if c.endswith('-2'):
                    item_level = 2
                elif c.endswith('-3'):
                    item_level = 3
            
            # Handle level changes
            while parents_stack and parents_stack[-1][0] >= item_level:
                parents_stack.pop()
            
            if not parents_stack:
                parents_stack = [(1, main_list)]
            
            parent_level, parent_list = parents_stack[-1]
            
            # Create new nested list if needed
            if item_level > parent_level:
                new_list = soup.new_tag('ul' if is_current_bullet else 'ol')
                if parent_list.contents:
                    parent_list.contents[-1].append(new_list)
                else:
                    temp_li = soup.new_tag('li')
                    temp_li.string = ''
                    parent_list.append(temp_li)
                    temp_li.append(new_list)
                parents_stack.append((item_level, new_list))
                parent_list = new_list
            
            # Create and append list item
            new_item = soup.new_tag('li')
            new_item.string = item.string
            parent_list.append(new_item)
            
        return main_list
    
    # Find and process all list items
    list_items = soup.find_all('li', class_=['list-bullet', 'list-bullet-2', 'list-bullet-3', 
                                            'list-number', 'list-number-2', 'list-number-3'])
    
    if list_items:
        # Create nested structure
        nested_list = nest_list_items(list_items)
        if nested_list:
            # Replace old items with new structure
            first_item = list_items[0]
            first_item.insert_before(nested_list)
            # Remove all original items
            for item in list_items:
                item.decompose()
    
    # Apply styles to lists
    for ul in soup.find_all('ul'):
        ul['style'] = 'list-style-type: disc; margin: 0.5em 0; padding-left: 2em;'
        
    for ol in soup.find_all('ol'):
        ol['style'] = 'list-style-type: decimal; margin: 0.5em 0; padding-left: 2em;'
    
    # Style nested lists
    for ul in soup.find_all('ul'):
        if ul.find_parent('li'):
            ul['style'] = 'list-style-type: circle; margin: 0.5em 0; padding-left: 2em;'
            for nested_ul in ul.find_all('ul'):
                nested_ul['style'] = 'list-style-type: square; margin: 0.5em 0; padding-left: 2em;'
    
    # Enhance tables
    for table in soup.find_all('table'):
        table['class'] = 'document-table'
        table['style'] = 'border-collapse: collapse; width: 100%; margin: 20px 0;'
        
        # Style table cells
        for cell in table.find_all(['td', 'th']):
            current_style = cell.get('style', '')
            cell['style'] = f'{current_style}; border: 1px solid #ddd; padding: 8px; text-align: left;'
        
        # Style header cells
        for header in table.find_all('th'):
            current_style = header.get('style', '')
            header['style'] = f'{current_style}; background-color: #f2f2f2; font-weight: bold;'
    
    # Enhance lists for better nesting
    for ul in soup.find_all('ul'):
        ul['style'] = 'margin: 10px 0; padding-left: 20px;'
        
    for ol in soup.find_all('ol'):
        ol['style'] = 'margin: 10px 0; padding-left: 20px;'
    
    for li in soup.find_all('li'):
        li['style'] = 'margin: 5px 0; line-height: 1.4;'
    
    # Handle nested lists better
    nested_lists = soup.find_all(['ul', 'ol'], recursive=True)
    for nested_list in nested_lists:
        if nested_list.parent.name == 'li':
            current_style = nested_list.get('style', '')
            nested_list['style'] = f'{current_style}; margin-top: 5px;'
    
    # Enhance paragraphs
    for p in soup.find_all('p'):
        current_style = p.get('style', '')
        p['style'] = f'{current_style}; margin: 10px 0; line-height: 1.6;'
    
    # Handle line breaks and preserve spacing
    for br in soup.find_all('br'):
        br['style'] = 'line-height: 1.6;'
    
    # Enhance images
    for img in soup.find_all('img'):
        current_style = img.get('style', '')
        img['style'] = f'{current_style}; max-width: 100%; height: auto; margin: 10px 0; display: block;'
        if not img.get('alt'):
            img['alt'] = 'Document image'
    
    return str(soup)

def create_complete_html(body_content):
    """
    Create a complete HTML document with enhanced CSS for complex content.
    """
    css = """
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }
        
        /* Table styles */
        .document-table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            font-size: 14px;
        }
        
        .document-table th,
        .document-table td {
            border: 1px solid #ddd;
            padding: 12px 8px;
            text-align: left;
            vertical-align: top;
        }
        
        .document-table th {
            background-color: #f8f9fa;
            font-weight: bold;
            color: #495057;
        }
        
        .document-table tbody tr:nth-child(even) {
            background-color: #f8f9fa;
        }
        
        .document-table tbody tr:hover {
            background-color: #e9ecef;
        }
        
        /* Multi-column layout support */
        .two-column {
            column-count: 2;
            column-gap: 20px;
            column-rule: 1px solid #ddd;
        }
        
        .three-column {
            column-count: 3;
            column-gap: 20px;
            column-rule: 1px solid #ddd;
        }
        
        /* List styles with better nesting */
        ul, ol {
            margin: 15px 0;
            padding-left: 30px;
        }
        
        li {
            margin: 8px 0;
            line-height: 1.5;
        }
        
        /* Nested list styles */
        ul ul, ol ol, ul ol, ol ul {
            margin: 5px 0;
            padding-left: 25px;
        }
        
        ul ul li {
            list-style-type: circle;
        }
        
        ul ul ul li {
            list-style-type: square;
        }
        
        /* Paragraph and spacing */
        p {
            margin: 12px 0;
            line-height: 1.6;
        }
        
        /* Headings */
        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            margin-top: 25px;
            margin-bottom: 15px;
            line-height: 1.3;
        }
        
        h1 { font-size: 2.2em; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
        h2 { font-size: 1.8em; color: #34495e; }
        h3 { font-size: 1.5em; color: #34495e; }
        h4 { font-size: 1.3em; color: #34495e; }
        h5 { font-size: 1.1em; color: #34495e; }
        h6 { font-size: 1em; color: #34495e; font-weight: bold; }
        
        /* Enhanced Image styles */
        img {
            max-width: 100% !important;
            height: auto !important;
            margin: 15px auto !important;
            display: block !important;
            border-radius: 4px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            background-color: #fff;
            object-fit: contain;
        }
        
        /* Image container for better positioning */
        .image-container {
            text-align: center;
            margin: 20px 0;
            max-width: 100%;
        }
        
        /* Handle responsive images */
        @media (max-width: 768px) {
            img {
                width: 100% !important;
                height: auto !important;
            }
        }
        
        /* Blockquotes */
        blockquote {
            border-left: 4px solid #3498db;
            margin: 20px 0;
            padding: 10px 20px;
            background-color: #ecf0f1;
            font-style: italic;
        }
        
        blockquote.intense {
            border-left-color: #e74c3c;
            background-color: #fadbd8;
        }
        
        /* Line breaks and spacing */
        br {
            line-height: 1.8;
        }
        
        /* Text formatting */
        strong, b {
            font-weight: 600;
            color: #2c3e50;
        }
        
        em, i {
            font-style: italic;
            color: #34495e;
        }
        
        /* Code blocks */
        pre, code {
            font-family: 'Courier New', Courier, monospace;
            background-color: #f8f9fa;
            padding: 2px 4px;
            border-radius: 3px;
            font-size: 0.9em;
        }
        
        pre {
            padding: 15px;
            border: 1px solid #e9ecef;
            border-radius: 4px;
            overflow-x: auto;
            margin: 15px 0;
        }
        
        /* Caption styles */
        .caption {
            font-size: 0.9em;
            color: #6c757d;
            font-style: italic;
            text-align: center;
            margin: 5px 0 15px 0;
        }
        
        /* Responsive design */
        @media (max-width: 768px) {
            body {
                padding: 10px;
            }
            
            .two-column, .three-column {
                column-count: 1;
            }
            
            .document-table {
                font-size: 12px;
            }
            
            .document-table th,
            .document-table td {
                padding: 6px 4px;
            }
        }
        
        /* Print styles */
        @media print {
            body {
                font-size: 12pt;
                line-height: 1.4;
            }
            
            .document-table {
                page-break-inside: avoid;
            }
            
            h1, h2, h3, h4, h5, h6 {
                page-break-after: avoid;
            }
        }
        
        /* Enhanced List styles */
        ul, ol {
            margin: 0.5em 0;
            padding-left: 2em;
            list-style-position: outside;
        }
        
        li {
            margin: 0.3em 0;
            line-height: 1.5;
            position: relative;
        }
        
        /* Nested list styles with proper indentation */
        li > ul,
        li > ol {
            margin: 0.5em 0;
            padding-left: 2em;
        }
        
        /* Bullet style hierarchy */
        ul {
            list-style-type: disc;
        }
        
        ul ul {
            list-style-type: circle;
        }
        
        ul ul ul {
            list-style-type: square;
        }
        
        /* Number style hierarchy */
        ol {
            list-style-type: decimal;
        }
        
        ol ol {
            list-style-type: lower-alpha;
        }
        
        ol ol ol {
            list-style-type: lower-roman;
        }
        
        /* Ensure proper spacing between nested items */
        li > ul:last-child,
        li > ol:last-child {
            margin-bottom: 0;
        }
    </style>
    """
    
    return f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Document Preview</title>
        {css}
    </head>
    <body>
        {body_content}
    </body>
    </html>
    """

def get_plain_text_from_html(html_content):
    """
    Extract plain text from HTML content by removing all HTML tags.
    
    Args:
        html_content (str): HTML content.
        
    Returns:
        str: Plain text content.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.decompose()
    
    # Get text and normalize whitespace
    text = soup.get_text()
    
    # Replace multiple whitespace with single space
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def analyze_document_structure(docx_path):
    """
    Analyze the document structure to provide insights about content complexity.
    
    Args:
        docx_path (str): Path to the Word document file.
        
    Returns:
        dict: Document analysis results.
    """
    try:
        doc = Document(docx_path)
        
        analysis = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
            'has_images': False,
            'styles_used': set(),
            'list_items': 0
        }
        
        # Analyze paragraphs
        for para in doc.paragraphs:
            if para.style.name:
                analysis['styles_used'].add(para.style.name)
            
            # Check for list items
            if 'List' in para.style.name:
                analysis['list_items'] += 1
        
        # Check for images by examining the document's relationship parts
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')
                for file_info in docx_zip.filelist:
                    if any(file_info.filename.lower().endswith(ext) for ext in image_extensions):
                        analysis['has_images'] = True
                        break
        except:
            pass
        
        analysis['styles_used'] = list(analysis['styles_used'])
        
        return analysis
        
    except Exception as e:
        return {'error': f'Could not analyze document: {str(e)}'}